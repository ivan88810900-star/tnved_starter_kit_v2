"""
Парсинг официальных PDF/DOCX «Таможенный тариф ЕАЭС» (pdfplumber / python-docx).

В выгрузках таблица часто не распознаётся как сетка — данные идут потоком текста;
используются строки с префиксами «РАЗДЕЛ», «ГРУППА» и строки с кодами ТН ВЭД.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import pdfplumber
from sqlalchemy import delete
from sqlalchemy.orm import Session

# Запуск из каталога backend: PYTHONPATH=.
_BACKEND_ROOT = Path(__file__).resolve().parent.parent.parent.parent  # .../backend

if str(_BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(_BACKEND_ROOT))

from app.db import Base, SessionLocal, engine  # noqa: E402
from app.models.tnved import Chapter, Commodity, Section  # noqa: E402

RE_SECTION = re.compile(r"^\s*РАЗДЕЛ\s+([IVXLCDM]+)\s*$", re.IGNORECASE)
RE_GROUP = re.compile(r"^\s*ГРУППА\s+(\d{2})\s*$", re.IGNORECASE)
# 10 знаков с пробелами: 0101 21 000 0 (группы 4 + три группы цифр)
RE_CODE_10 = re.compile(r"^\s*(\d{4}(?:\s+\d+){3})\s+(.*)$")
# 4 знака в начале строки (товарная позиция); далее — наименование (не подсубпозиция)
RE_CODE_4 = re.compile(r"^\s*(\d{4})\s+([А-Яа-яЁёA-Za-z«»].*)$")
RE_TABLE_HEADER = re.compile(r"Код.*ТН\s*ВЭД|Ставка\s+ввозной", re.IGNORECASE)

DATA_DIR = Path(__file__).resolve().parent / "data"
SUPPORTED_SUFFIXES = {".pdf", ".docx"}


def normalize_hs_code(raw: str) -> str:
    """Удаляет пробелы и переносы, оставляет только цифры."""
    return re.sub(r"\D", "", raw or "")


def chapter_num_to_roman(ch: int) -> str:
    """Номер главы ТН ВЭД (01–97) → римский номер раздела (структура ЕАЭС)."""
    if ch < 1 or ch > 97:
        return "I"
    if 1 <= ch <= 5:
        return "I"
    if 6 <= ch <= 14:
        return "II"
    if ch == 15:
        return "III"
    if 16 <= ch <= 24:
        return "IV"
    if 25 <= ch <= 27:
        return "V"
    if 28 <= ch <= 31:
        return "VI"
    if ch == 32:
        return "VII"
    if 33 <= ch <= 38:
        return "VIII"
    if 39 <= ch <= 40:
        return "IX"
    if 41 <= ch <= 43:
        return "X"
    if 44 <= ch <= 46:
        return "XI"
    if 47 <= ch <= 49:
        return "XII"
    if 50 <= ch <= 63:
        return "XIII"
    if 64 <= ch <= 67:
        return "XIV"
    if 68 <= ch <= 70:
        return "XV"
    if ch == 71:
        return "XVI"
    if 72 <= ch <= 83:
        return "XVII"
    if 84 <= ch <= 85:
        return "XVIII"
    if 86 <= ch <= 89:
        return "XIX"
    if 90 <= ch <= 92:
        return "XX"
    return "XXI"


def _parse_commodity_tail(rest: str) -> tuple[str, str, str]:
    """Описание, ед. изм., ставка пошлины (текст).

    В PDF часто:
    - «… – 50, но не менее» (ставка не в конце строки);
    - «… шт 563С)» (OCR: буква вместо % или лишняя скобка);
    - только тире «–» как льгота/0 %.
    """
    rest = rest.strip()
    if not rest:
        return "", "", ""

    # 1) «… – 50, но не менее …» или «50 %, но не менее»
    m = re.search(
        r"[–\-]\s*(\d+(?:[.,]\d+)?)\s*%?\s*,?\s*(но\s+не\s+менее[^\n.]*)",
        rest,
        re.IGNORECASE,
    )
    if m:
        duty = rest[m.start() :].strip()
        desc = rest[: m.start()].strip()
        return desc, "", duty

    # 1b) «… – 1053С)», «… – 52С)» (ставка без %; в PDF «С» + «)» вместо %)
    m = re.search(r"[–\-]\s*(\d+(?:[.,]\d+)?)\s*[СC]\)\s*$", rest)
    if m:
        duty = rest[m.start() :].strip()
        desc = rest[: m.start()].strip()
        return desc, "", duty

    # 1c) «… шт 5, но не менее» (без тире непосредственно перед цифрой)
    m = re.search(
        r"\s+(шт|кг|л|1000\s*шт|тыс\.\s*шт|м2|м3|m²|m2|п\.м\.)\s+(\d+(?:[.,]\d+)?)\s*,?\s*(но\s+не\s+менее[^\n.]*)",
        rest,
        re.IGNORECASE,
    )
    if m:
        duty = rest[m.start() :].strip()
        desc = rest[: m.start()].strip()
        return desc, "", duty

    # 1d) «… – 0,06 евро за 1 кг», «– 0,015 евро»
    m = re.search(r"[–\-]\s*(\d+(?:[.,]\d+)?)\s*евро[^\n]*$", rest, re.IGNORECASE)
    if m:
        duty = rest[m.start() :].strip()
        desc = rest[: m.start()].strip()
        return desc, "", duty

    # 1e) «… – 250 долларов», «– 171 доллар США»
    m = re.search(
        r"[–\-]\s*(\d+(?:[.,]\d+)?)\s*(?:долларов|доллара(?:\s+США)?|доллар(?:\s+США)?)[^\n]*$",
        rest,
        re.IGNORECASE,
    )
    if m:
        duty = rest[m.start() :].strip()
        desc = rest[: m.start()].strip()
        return desc, "", duty

    # 2) Ед. изм. + число + % и хвост OCR («563С)», «5 %»)
    m = re.search(
        r"\s+(шт|кг|л|1000\s*шт|тыс\.\s*шт|м2|м3|m²|m2|п\.м\.)\s+(\d+(?:[.,]\d+)?)\s*%?\s*[%СCс)\]°º]*\s*$",
        rest,
        re.IGNORECASE,
    )
    if m:
        desc = rest[: m.start()].strip()
        return desc, m.group(1), m.group(2)

    # 3) Только число / процент в конце строки
    m2 = re.search(r"\s+(\d+(?:[.,]\d+)?%?)\s*$", rest)
    if m2:
        return rest[: m2.start()].strip(), "", m2.group(1)

    # 5) Ставка «0» или «–» как отдельный маркер в конце (после длинного тире)
    m3 = re.search(r"[–\-]\s*(0|–|—)\s*$", rest)
    if m3:
        return rest[: m3.start()].strip(), "", m3.group(1)

    return rest, "", ""


def _iter_lines_pdf(pdf_path: Path):
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_idx, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            for line in text.splitlines():
                yield page_idx, line


def _iter_lines_docx(docx_path: Path):
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("Для .docx установите: pip install python-docx") from e

    doc = Document(str(docx_path))
    pi = 0
    for para in doc.paragraphs:
        pi += 1
        for line in (para.text or "").splitlines():
            if line.strip():
                yield pi, line
    for ti, table in enumerate(doc.tables, start=1):
        for row in table.rows:
            line = " ".join((c.text or "").strip() for c in row.cells)
            if line.strip():
                yield ti, line


def iter_lines_from_path(path: Path):
    """Единый поток (номер_фрагмента, строка) для PDF и DOCX."""
    suf = path.suffix.lower()
    if suf == ".pdf":
        yield from _iter_lines_pdf(path)
    elif suf == ".docx":
        yield from _iter_lines_docx(path)
    else:
        raise ValueError(f"Неподдерживаемый формат: {path} (ожидались {SUPPORTED_SUFFIXES})")


def _flush_table_lines(lines: list[tuple[int, str]]) -> list[tuple[int, str]]:
    out: list[tuple[int, str]] = []
    for p, line in lines:
        if RE_TABLE_HEADER.search(line) and ("Код" in line or "Ставка" in line):
            continue
        out.append((p, line))
    return out


def discover_tariff_files(data_dir: Path) -> list[Path]:
    """Все PDF и DOCX в каталоге, по алфавиту (ru.01… перед ru.02…)."""
    if not data_dir.is_dir():
        return []
    found: list[Path] = []
    for p in data_dir.iterdir():
        if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES:
            found.append(p)
    return sorted(found, key=lambda x: x.name.lower())


def parse_tariff_rows(
    session: Session,
    rows: list[tuple[int, str]],
    *,
    verbose: bool = False,
) -> tuple[int, int, int]:
    """
    Разбор уже извлечённых строк одного файла.
    Раздел/группа: при совпадении в БД — переиспользование (без дублирования записей).
    Возвращает (новых разделов, новых групп, новых позиций).
    """
    rows = _flush_table_lines(rows)
    n_sections = n_chapters = n_commodities = 0
    section: Section | None = None
    chapter: Chapter | None = None
    # Дубликаты в одном файле до commit не видны в БД — учитываем в памяти
    seen_commodity: set[tuple[int, str]] = set()
    i = 0
    n = len(rows)

    def _log(msg: str) -> None:
        if verbose:
            print(msg)

    while i < n:
        page_no, line = rows[i]
        s = line.strip()
        if not s:
            i += 1
            continue

        m_sec = RE_SECTION.match(s)
        if m_sec:
            roman = m_sec.group(1).upper()
            existing = session.query(Section).filter(Section.roman_number == roman).first()
            if existing is not None:
                section = existing
                chapter = None
                _log(f"[parse] фрагмент {page_no}: РАЗДЕЛ {roman} — уже в БД (id={section.id}), пропуск шапки…")
                i += 1
                while i < n:
                    _, ln = rows[i]
                    t = ln.strip()
                    if RE_GROUP.match(t) or RE_SECTION.match(t):
                        break
                    i += 1
                continue

            title = ""
            notes_lines: list[str] = []
            i += 1
            if i < n:
                _, title_line = rows[i]
                title = title_line.strip()
                i += 1
            _log(f"[parse] фрагмент {page_no}: РАЗДЕЛ {roman} — {title[:80]!r}")

            while i < n:
                _, ln = rows[i]
                t = ln.strip()
                if RE_GROUP.match(t):
                    break
                if RE_SECTION.match(t):
                    break
                notes_lines.append(ln.rstrip())
                i += 1

            section = Section(
                roman_number=roman,
                title=title,
                notes="\n".join(notes_lines).strip(),
            )
            session.add(section)
            session.flush()
            n_sections += 1
            chapter = None
            continue

        m_grp = RE_GROUP.match(s)
        if m_grp:
            if section is None:
                gc = m_grp.group(1)
                roman = chapter_num_to_roman(int(gc))
                section = session.query(Section).filter(Section.roman_number == roman).first()
                if section is None:
                    section = Section(roman_number=roman, title="", notes="")
                    session.add(section)
                    session.flush()
                    n_sections += 1
                chapter = None

            code = m_grp.group(1)
            existing_ch = (
                session.query(Chapter)
                .filter(Chapter.section_id == section.id, Chapter.code == code)
                .first()
            )
            if existing_ch is not None:
                chapter = existing_ch
                _log(f"[parse] фрагмент {page_no}: ГРУППА {code} — уже в БД (id={chapter.id}), пропуск шапки…")
                i += 1
                while i < n:
                    _, ln = rows[i]
                    t = ln.strip()
                    if RE_GROUP.match(t) or RE_SECTION.match(t):
                        break
                    if RE_CODE_10.match(t) or RE_CODE_4.match(t):
                        break
                    i += 1
                continue

            title_ch = ""
            notes_ch: list[str] = []
            i += 1
            if i < n:
                _, title_line = rows[i]
                title_ch = title_line.strip()
                i += 1
            _log(f"[parse] фрагмент {page_no}: ГРУППА {code} — {title_ch[:80]!r}")

            while i < n:
                _, ln = rows[i]
                t = ln.strip()
                if RE_GROUP.match(t) or RE_SECTION.match(t):
                    break
                if RE_CODE_10.match(t) or RE_CODE_4.match(t):
                    break
                notes_ch.append(ln.rstrip())
                i += 1

            chapter = Chapter(
                section_id=section.id,
                code=code,
                title=title_ch,
                notes="\n".join(notes_ch).strip(),
            )
            session.add(chapter)
            session.flush()
            n_chapters += 1
            continue

        if section is not None and chapter is not None:
            m10 = RE_CODE_10.match(s)
            if m10:
                raw_code = m10.group(1)
                rest = m10.group(2)
                code_norm = normalize_hs_code(raw_code)
                ck = (chapter.id, code_norm)
                if ck in seen_commodity:
                    _log(f"[parse] фрагмент {page_no}: пропуск дубликата кода {code_norm} (тот же файл)")
                    i += 1
                    continue
                dup = (
                    session.query(Commodity)
                    .filter(Commodity.chapter_id == chapter.id, Commodity.code == code_norm)
                    .first()
                )
                if dup is not None:
                    seen_commodity.add(ck)
                    _log(f"[parse] фрагмент {page_no}: пропуск дубликата кода {code_norm} в БД")
                    i += 1
                    continue
                desc, unit, duty = _parse_commodity_tail(rest)
                _log(f"[parse] фрагмент {page_no}: код {code_norm} ({len(code_norm)} зн.) — {desc[:60]!r}…")
                session.add(
                    Commodity(
                        chapter_id=chapter.id,
                        code=code_norm,
                        description=desc,
                        unit=unit,
                        import_duty=duty,
                    )
                )
                seen_commodity.add(ck)
                n_commodities += 1
                i += 1
                continue

            m4 = RE_CODE_4.match(s)
            if m4:
                raw_code = m4.group(1)
                rest = m4.group(2)
                code_norm = normalize_hs_code(raw_code)
                ck = (chapter.id, code_norm)
                if ck in seen_commodity:
                    _log(f"[parse] фрагмент {page_no}: пропуск дубликата кода {code_norm} (тот же файл)")
                    i += 1
                    continue
                dup = (
                    session.query(Commodity)
                    .filter(Commodity.chapter_id == chapter.id, Commodity.code == code_norm)
                    .first()
                )
                if dup is not None:
                    seen_commodity.add(ck)
                    _log(f"[parse] фрагмент {page_no}: пропуск дубликата кода {code_norm} в БД")
                    i += 1
                    continue
                desc, unit, duty = _parse_commodity_tail(rest)
                _log(f"[parse] фрагмент {page_no}: код {code_norm} ({len(code_norm)} зн.) — {desc[:60]!r}…")
                session.add(
                    Commodity(
                        chapter_id=chapter.id,
                        code=code_norm,
                        description=desc,
                        unit=unit,
                        import_duty=duty,
                    )
                )
                seen_commodity.add(ck)
                n_commodities += 1
                i += 1
                continue

        i += 1

    return n_sections, n_chapters, n_commodities


def parse_file_to_db(
    session: Session,
    file_path: Path,
    *,
    verbose: bool = False,
) -> tuple[int, int, int]:
    """Извлекает текст из PDF/DOCX и передаёт в parse_tariff_rows."""
    if not file_path.is_file():
        raise FileNotFoundError(f"Файл не найден: {file_path}")
    rows = list(iter_lines_from_path(file_path))
    return parse_tariff_rows(session, rows, verbose=verbose)


def clear_tnved_tables(session: Session) -> None:
    session.execute(delete(Commodity))
    session.execute(delete(Chapter))
    session.execute(delete(Section))
    session.commit()


def extract_tables_debug(pdf_path: Path, max_pages: int = 3) -> None:
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_idx, page in enumerate(pdf.pages[:max_pages], start=1):
            tables = page.extract_tables() or []
            print(f"[tables] страница {page_idx}: таблиц: {len(tables)}")
            for ti, table in enumerate(tables):
                print(f"  table {ti}: {len(table)} строк")
                for row in table[:5]:
                    print(f"    {row}")


def _batch_log(msg: str, *, use_tqdm: bool) -> None:
    if use_tqdm:
        try:
            from tqdm import tqdm as tqdm_cls

            tqdm_cls.write(msg)
        except Exception:
            print(msg)
    else:
        print(msg)


def run_batch(
    data_dir: Path,
    *,
    clear_first: bool = False,
    verbose_lines: bool = False,
    single_file: Path | None = None,
    use_tqdm: bool = True,
    continue_on_error: bool = False,
) -> int:
    """
    Обработка всех PDF/DOCX в data/ с tqdm и commit после каждого файла.
    Возвращает код выхода: 0 — успех, 1 — была ошибка (если continue_on_error).
    """
    tqdm_cls = None
    if use_tqdm:
        try:
            from tqdm import tqdm as tqdm_cls
        except ImportError as e:
            raise ImportError("Установите: pip install tqdm") from e

    print()
    print("=" * 60)
    print("  Таможенный тариф ЕАЭС — каталог → SQLite (tnved.db)")
    print("=" * 60)
    print(f"[batch] База: {engine.url}")
    print(f"[batch] Каталог данных: {data_dir.resolve()}")

    Base.metadata.create_all(bind=engine)

    files: list[Path]
    if single_file is not None:
        if not single_file.is_file():
            print(f"[batch] ОШИБКА: файл не найден: {single_file}")
            sys.exit(1)
        files = [single_file]
    else:
        files = discover_tariff_files(data_dir)

    if not files:
        print(f"[batch] Нет файлов {SUPPORTED_SUFFIXES} в {data_dir}")
        sys.exit(1)

    print(f"[batch] К обработке: {len(files)} файл(ов) (порядок: по имени A→Z)")

    db = SessionLocal()
    totals = {"sections": 0, "chapters": 0, "commodities": 0}
    had_error = False
    try:
        if clear_first:
            clear_tnved_tables(db)
            print("[batch] Очищены таблицы tnved_sections, tnved_chapters, tnved_commodities.")

        if use_tqdm and tqdm_cls is not None:
            file_iter = tqdm_cls(
                files,
                desc="Обработка файлов",
                unit="файл",
                ncols=110,
                bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}] {postfix}",
            )
        else:
            file_iter = files

        for idx, path in enumerate(file_iter, start=1):
            if use_tqdm and tqdm_cls is not None:
                file_iter.set_postfix_str(path.name[:36] + ("…" if len(path.name) > 36 else ""))
            _batch_log(f"\n[batch] Обработка файла {idx} из {len(files)}: {path.name}", use_tqdm=use_tqdm)
            try:
                ns, nc, nx = parse_file_to_db(db, path, verbose=verbose_lines)
                db.commit()
                totals["sections"] += ns
                totals["chapters"] += nc
                totals["commodities"] += nx
                _batch_log(
                    f"    ✓ OK: +разделов {ns}, +групп {nc}, +позиций {nx} (commit выполнен)",
                    use_tqdm=use_tqdm,
                )
            except Exception as ex:
                db.rollback()
                had_error = True
                _batch_log(
                    f"    ✗ ОШИБКА: {ex} — rollback этого файла; предыдущие commit сохранены",
                    use_tqdm=use_tqdm,
                )
                if not continue_on_error:
                    raise
    finally:
        db.close()

    print()
    print("[batch] Итого добавлено за этот запуск (новые строки):")
    print(
        f"       разделов: {totals['sections']}, групп: {totals['chapters']}, "
        f"позиций: {totals['commodities']}"
    )

    try:
        import pandas as pd

        with engine.connect() as conn:
            df = pd.read_sql_query(
                "SELECT COUNT(*) AS n FROM tnved_commodities",
                conn,
            )
        print("\n[batch] Всего строк в tnved_commodities:", int(df["n"].iloc[0]))
    except Exception as ex:
        print(f"[batch] pandas summary: {ex}")

    return 1 if had_error else 0


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Импорт ТН ВЭД из PDF/DOCX в SQLite")
    parser.add_argument(
        "path",
        nargs="?",
        default=None,
        help="Один PDF/DOCX или не указывать — все файлы из data/",
    )
    parser.add_argument(
        "--clear",
        action="store_true",
        help="Перед импортом очистить tnved_* таблицы",
    )
    parser.add_argument(
        "--verbose-lines",
        action="store_true",
        help="Подробные строки парсера (много вывода)",
    )
    parser.add_argument(
        "--no-progress",
        action="store_true",
        help="Без tqdm (только текст «файл N из M»)",
    )
    parser.add_argument(
        "--continue-on-error",
        action="store_true",
        help="После ошибки в файле продолжить со следующим (по умолчанию — остановка)",
    )
    args = parser.parse_args()

    data_dir = DATA_DIR
    single: Path | None = None
    if args.path:
        p = Path(args.path).resolve()
        if p.is_dir():
            data_dir = p
        else:
            single = p

    code = run_batch(
        data_dir,
        clear_first=args.clear,
        verbose_lines=args.verbose_lines,
        single_file=single,
        use_tqdm=not args.no_progress,
        continue_on_error=args.continue_on_error,
    )
    sys.exit(code)
