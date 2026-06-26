"""Массовый ИИ-импорт нормативных документов (PDF/DOCX/HTML) с лимитами API и чекпоинтами."""

from __future__ import annotations

import asyncio
import hashlib
import io
import json
import os
import re
import time
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Callable

from loguru import logger
from sqlalchemy.orm import Session

from ..datetime_util import utc_now_naive
from ..db import SessionLocal
from ..models import BulkImportFileCheckpoint, BulkImportJob, HsRate
from ..models.tnved import Commodity, NonTariffMeasure, SpecialDuty
from .gemini_genai_configure import normalize_gemini_api_endpoint_for_sdk, resolved_gemini_model_name
from .normative_store import normalize_hs_duty_rate_string

BACKEND_ROOT = Path(__file__).resolve().parents[2]
RAW_NORMATIVE_DIR = BACKEND_ROOT / "data" / "raw_normative"
LLM_RAW_DIR = BACKEND_ROOT / "logs" / "llm_raw"

BULK_SYSTEM_PROMPT = (
    "Ты анализируешь таможенный нормативный акт. Твоя задача — извлечь конкретные меры регулирования. "
    "Определи: 1) Тип меры (Акциз, Нетарифка, Запрет на ввоз, Пошлина), 2) Коды ТН ВЭД "
    "(могут быть указаны префиксы, например «из 8516»), 3) Ставку или требование "
    "(например «20%» или «Лицензия ФСТЭК»). Верни массив в формате JSON.\n\n"
    "Критически важно различать ввозную таможенную пошлину (duty), акциз (excise) и НДС (VAT). "
    "Не путай НДС со ставкой ввозной пошлины: для пошлины используй measure_category «duty» и ставку в "
    "rate_or_requirement; для НДС указывай отдельное поле vat_rate. "
    "Если в тексте указана дата вступления в силу нормы или акта, обязательно укажи её в valid_from "
    "(формат YYYY-MM-DD).\n\n"
    "Обязательно нормализуй значения ставок в rate_or_requirement (для пошлины/акциза) и в vat_rate. "
    "Превращай текстовые описания в числа со знаком процента: «0 процентов» или «ноль процентов» — строго в "
    "\"0%\"; «десять процентов» — в \"10%\"; убирай лишние слова (равна, составляет и т.п.). "
    "Числовое значение в JSON для vat_rate допускается только если оно уже в процентах (например 20).\n\n"
    "Каждый объект массива используй СТРОГО с английскими ключами:\n"
    '- "measure_category": одно из: excise, non_tariff, import_ban, duty, special_duty, other\n'
    '- "hs_codes": массив строк (только цифры, длина 4–10 символов; допускаются префиксы)\n'
    '- "rate_or_requirement": краткая строка с формулировкой ставки или требования\n'
    '- "regulatory_act": наименование или реквизиты акта\n'
    '- "origin_country": ISO-3166-1 alpha-2 для спецпошлины, иначе пустая строка\n'
    '- "valid_from": строка даты вступления в силу в формате YYYY-MM-DD (если нет в тексте — пустая строка)\n'
    '- "vat_rate": ставка НДС при импорте в процентах или как число (например 20 или \"20%\"); '
    "если не применимо — пустая строка или опусти ключ\n\n"
    "Верни только валидный JSON-массив без markdown и без пояснений."
)

_MAX_TEXT_CHARS = 48_000
_last_llm_monotonic: float = 0.0
_import_lock = asyncio.Lock()
_job_running = False

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".html", ".htm"}


def raw_normative_dir() -> Path:
    RAW_NORMATIVE_DIR.mkdir(parents=True, exist_ok=True)
    return RAW_NORMATIVE_DIR


def _sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def list_input_files(directory: Path | None = None) -> list[Path]:
    root = directory or raw_normative_dir()
    if not root.is_dir():
        return []
    out: list[Path] = []
    for p in sorted(root.rglob("*")):
        if not p.is_file():
            continue
        if p.name.startswith("."):
            continue
        if p.suffix.lower() in SUPPORTED_SUFFIXES:
            out.append(p)
    return out


def _extract_pdf_text(path: Path) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    parts: list[str] = []
    try:
        for page in doc:
            parts.append(page.get_text("text") or "")
    finally:
        doc.close()
    return "\n".join(parts).strip()


def _extract_docx_text(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    return "\n".join((p.text or "").strip() for p in document.paragraphs if (p.text or "").strip()).strip()


def _extract_html_text(path: Path) -> str:
    from bs4 import BeautifulSoup

    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return (soup.get_text("\n", strip=True) or "").strip()


def extract_text_from_file(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".pdf":
        return _extract_pdf_text(path)
    if suf == ".docx":
        return _extract_docx_text(path)
    if suf in (".html", ".htm"):
        return _extract_html_text(path)
    raise ValueError(f"Неподдерживаемый формат: {path.suffix}")


def html_string_to_text(html: str) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html or "", "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return (soup.get_text("\n", strip=True) or "").strip()


def _extract_pdf_bytes(data: bytes) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype="pdf")
    parts: list[str] = []
    try:
        for page in doc:
            parts.append(page.get_text("text") or "")
    finally:
        doc.close()
    return "\n".join(parts).strip()


def _extract_docx_bytes(data: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    return "\n".join((p.text or "").strip() for p in document.paragraphs if (p.text or "").strip()).strip()


def extract_text_from_bytes(data: bytes, source_hint: str = "") -> str:
    """Извлечь текст из PDF / DOCX / HTML по сигнатурам и подсказке URL."""
    if not data:
        return ""
    if data[:4] == b"%PDF":
        try:
            return _extract_pdf_bytes(data)
        except Exception:
            return ""
    if len(data) > 4 and data[:2] == b"PK":
        try:
            return _extract_docx_bytes(data)
        except Exception:
            pass
    hint = (source_hint or "").lower()
    if any(hint.endswith(x) for x in (".pdf",)) or (b"%PDF" in data[:2000]):
        try:
            return _extract_pdf_bytes(data)
        except Exception:
            pass
    if any(hint.endswith(x) for x in (".docx", ".doc")):
        try:
            return _extract_docx_bytes(data)
        except Exception:
            pass
    for enc in ("utf-8", "cp1251"):
        try:
            raw = data.decode(enc, errors="replace")
            txt = html_string_to_text(raw)
            if len(txt) > 120:
                return txt
        except Exception:
            continue
    return html_string_to_text(data.decode("utf-8", errors="replace"))


def _extract_json_array(raw: str) -> str:
    text = (raw or "").strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)
    if text.startswith("[") and text.endswith("]"):
        return text
    i = text.find("[")
    j = text.rfind("]")
    if i != -1 and j != -1 and j > i:
        return text[i : j + 1]
    raise ValueError("В ответе модели не найден JSON-массив")


def _is_rate_limit_error(exc: BaseException) -> bool:
    try:
        from google.api_core import exceptions as ga_exc

        if isinstance(exc, ga_exc.ResourceExhausted):
            return True
    except Exception:
        pass
    s = str(exc).lower()
    return (
        "429" in s
        or "resource exhausted" in s
        or "too many requests" in s
        or "quota" in s
        or "rate limit" in s
    )


def _normalize_hs_token(raw: str) -> str:
    d = re.sub(r"\D", "", raw or "")
    if len(d) < 4:
        return ""
    return d[:10]


def _normalize_measure_category(raw: object) -> str:
    s = str(raw or "").strip().lower()
    mapping = {
        "акциз": "excise",
        "excise": "excise",
        "нетариф": "non_tariff",
        "non_tariff": "non_tariff",
        "запрет": "import_ban",
        "import_ban": "import_ban",
        "пошлин": "duty",
        "duty": "duty",
        "спец": "special_duty",
        "special_duty": "special_duty",
        "special": "special_duty",
    }
    for needle, cat in mapping.items():
        if needle in s:
            return cat
    if s in ("excise", "non_tariff", "import_ban", "duty", "special_duty", "other"):
        return s
    return "other"


def _nt_measure_type_for_non_tariff(text: str) -> str:
    t = (text or "").lower()
    if "запрет" in t or "embargo" in t:
        return "ban"
    if "вет" in t or "ветерин" in t:
        return "vet_control"
    if "фито" in t:
        return "phyto_control"
    if "тр тс" in t or "техреглам" in t or " tr " in t:
        return "tr_ts"
    if "лиценз" in t or "фстек" in t or "фсб" in t or "фсвтс" in t:
        return "license"
    return "certificate"


_RU_PERCENT_WORDS = re.compile(
    r"(?iu)(?:процентов|процента|процент)"
)


def _normalize_rate_string_for_percent(value: str | Any) -> str:
    """Убрать русские слова «процент*», нормализовать запятую в точку для парсинга ставки."""
    if value is None:
        return ""
    s = str(value).strip().replace("\xa0", " ")
    if not s:
        return ""
    # Частые текстовые нули/числа в актах (после удаления «процент*» останется цифра)
    s = re.sub(r"(?iu)\b(ноль|нуль)\b", "0", s)
    s = re.sub(r"(?iu)\bдесять\b", "10", s)
    s = _RU_PERCENT_WORDS.sub(" ", s)
    s = s.replace(",", ".")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _parse_percent(value: str | Any) -> float | None:
    """Доля в процентах из строки (в т.ч. «0 процентов», «15 %»)."""
    s = _normalize_rate_string_for_percent(value)
    if not s:
        return None
    m = re.search(r"(\d+(?:\.\d+)?)\s*%", s, re.IGNORECASE)
    if m:
        try:
            v = float(m.group(1))
            return v if 0.0 <= v <= 100.0 else None
        except ValueError:
            return None
    m2 = re.fullmatch(r"\D*(\d+(?:\.\d+)?)\D*", s)
    if m2:
        try:
            v = float(m2.group(1))
            return v if 0.0 <= v <= 100.0 else None
        except ValueError:
            return None
    return None


def _parse_fixed_rub(value: str) -> float | None:
    m = re.search(r"(\d+(?:[.,]\d+)?)\s*(?:руб|руб\.|₽)", value or "", re.IGNORECASE)
    if not m:
        return None
    try:
        return float(m.group(1).replace(",", "."))
    except ValueError:
        return None


def _resolve_commodity_codes(db: Session, hs_codes: list[str], *, max_codes: int = 400) -> list[str]:
    found: list[str] = []
    seen: set[str] = set()
    for raw in hs_codes:
        d = _normalize_hs_token(raw)
        if not d:
            continue
        if len(d) == 10:
            row = db.query(Commodity.code).filter(Commodity.code == d).first()
            if row and row[0] not in seen:
                found.append(row[0])
                seen.add(row[0])
            continue
        q = db.query(Commodity.code).filter(Commodity.code.like(f"{d}%"))
        for (code,) in q.all():
            if len(code) == 10 and code not in seen:
                found.append(code)
                seen.add(code)
            if len(found) >= max_codes:
                return found
    return found


def _hs_prefix_for_rates(digits: str) -> str:
    d = re.sub(r"\D", "", digits or "")
    if len(d) >= 10:
        return d[:10]
    if len(d) >= 6:
        return d[:6]
    if len(d) >= 4:
        return d[:4]
    return d[:10] if d else ""


_VALID_FROM_RE = re.compile(r"^(19|20)\d{2}-(0[1-9]|1[0-2])-(0[1-9]|[12]\d|3[01])$")


def _parse_vat_rate_value(value: Any) -> float | None:
    """Ставка НДС в процентах для hs_rates.vat_import_rate."""
    if value is None:
        return None
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        v = float(value)
        return v if 0.0 <= v <= 100.0 else None
    s = _normalize_rate_string_for_percent(str(value).strip())
    if not s:
        return None
    p = _parse_percent(s)
    if p is not None:
        return p
    try:
        v = float(s.rstrip("%").strip())
    except ValueError:
        return None
    return v if 0.0 <= v <= 100.0 else None


def _parse_valid_from_str(value: Any) -> str:
    """Дата вступления в силу для hs_rates.valid_from (YYYY-MM-DD)."""
    if value is None:
        return ""
    s = str(value).strip()[:32]
    if not s:
        return ""
    m = re.search(r"(19|20)\d{2}-(0[1-9]|1[0-2])-(0[1-9]|[12]\d|3[01])", s)
    if not m:
        return ""
    cand = m.group(0)
    if not _VALID_FROM_RE.match(cand):
        return ""
    y, mo, d = (int(cand[0:4]), int(cand[5:7]), int(cand[8:10]))
    try:
        date(y, mo, d)
    except ValueError:
        return ""
    return cand


def _apply_hs_rate_fields_from_row(obj: HsRate, row: dict[str, Any]) -> None:
    """Доп. поля из JSON: vat_rate → vat_import_rate, valid_from → valid_from."""
    vat_f = _parse_vat_rate_value(row.get("vat_rate"))
    if vat_f is not None:
        obj.vat_import_rate = float(vat_f)
    vf = _parse_valid_from_str(row.get("valid_from"))
    if vf:
        obj.valid_from = vf[:20]


def _save_llm_raw_response(raw: str) -> None:
    """Сохранить сырой текст ответа модели для отладки (в т.ч. при measures_applied=0)."""
    try:
        LLM_RAW_DIR.mkdir(parents=True, exist_ok=True)
        ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        path = LLM_RAW_DIR / f"doc_{ts}.json"
        payload = {"saved_at": datetime.now(timezone.utc).isoformat(), "raw_text": raw or ""}
        path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
        logger.debug("llm_raw: сохранён ответ в {}", path)
    except Exception as e:
        logger.warning("llm_raw: не удалось сохранить ответ модели: {}", e)


def apply_structured_rows(
    db: Session,
    rows: list[dict[str, Any]],
    *,
    source_tag: str,
) -> int:
    """UPSERT в non_tariff_measures / special_duties / hs_rates (акцизы и пошлины). Возвращает число применённых строк."""
    applied = 0
    rev = f"bulk-ai:{source_tag[:80]}"

    for row in rows:
        if not isinstance(row, dict):
            continue
        cat = _normalize_measure_category(row.get("measure_category"))
        hs_codes = row.get("hs_codes") if isinstance(row.get("hs_codes"), list) else []
        hs_strs = [str(x) for x in hs_codes if x is not None]
        rate_txt = str(row.get("rate_or_requirement") or "").strip()
        act = str(row.get("regulatory_act") or "").strip()[:255] or "Импорт из нормативного архива"
        origin = str(row.get("origin_country") or "").strip().upper()[:8]

        if cat == "import_ban":
            measure_type = "ban"
        elif cat == "non_tariff":
            measure_type = _nt_measure_type_for_non_tariff(rate_txt + " " + act)
        else:
            measure_type = ""

        codes = _resolve_commodity_codes(db, hs_strs)
        prefix_union = ""
        if hs_strs:
            prefix_union = _hs_prefix_for_rates(_normalize_hs_token(hs_strs[0]) or hs_strs[0])

        if cat in ("import_ban", "non_tariff") and codes:
            doc_req = rate_txt[:255] if len(rate_txt) <= 255 else rate_txt[:252] + "…"
            desc = (rate_txt + " — " + act)[:4000]
            for c in codes:
                existing = (
                    db.query(NonTariffMeasure)
                    .filter(
                        NonTariffMeasure.commodity_code == c,
                        NonTariffMeasure.measure_type == measure_type,
                        NonTariffMeasure.regulatory_act == act,
                    )
                    .first()
                )
                if existing:
                    existing.description = desc[:4000]
                    existing.document_required = doc_req
                else:
                    db.add(
                        NonTariffMeasure(
                            commodity_code=c,
                            measure_type=measure_type,
                            description=desc[:4000],
                            document_required=doc_req,
                            regulatory_act=act,
                        )
                    )
                applied += 1

        elif cat == "special_duty" and prefix_union and origin and (rate_txt or "").strip():
            pct = _parse_percent(rate_txt) or 0.0
            spec = _parse_fixed_rub(rate_txt) or 0.0
            if pct <= 0 and spec <= 0:
                continue
            existing = (
                db.query(SpecialDuty)
                .filter(
                    SpecialDuty.hs_code_prefix == prefix_union[:16],
                    SpecialDuty.origin_country == origin,
                    SpecialDuty.regulatory_act == act,
                )
                .first()
            )
            if existing:
                existing.rate_percent = float(pct)
                existing.rate_specific = float(spec)
                existing.currency_code = "RUB"
            else:
                db.add(
                    SpecialDuty(
                        hs_code_prefix=prefix_union[:16],
                        origin_country=origin,
                        rate_percent=float(pct),
                        rate_specific=float(spec),
                        currency_code="RUB",
                        regulatory_act=act,
                    )
                )
            applied += 1

        elif cat == "duty" and prefix_union:
            pct = _parse_percent(rate_txt)
            vat_f = _parse_vat_rate_value(row.get("vat_rate"))
            vf = _parse_valid_from_str(row.get("valid_from"))
            if pct is None and vat_f is None and not vf:
                continue
            pref = prefix_union[:10]
            rt = (rate_txt or "").strip()
            duty_stored: str | None
            if rt:
                duty_stored = normalize_hs_duty_rate_string(rt)
            elif pct is not None:
                duty_stored = normalize_hs_duty_rate_string(pct)
            else:
                duty_stored = None
            obj = db.query(HsRate).filter(HsRate.hs_prefix == pref).first()
            if obj:
                if duty_stored is not None:
                    obj.duty_rate = duty_stored
                obj.source_revision = rev[:128]
                _apply_hs_rate_fields_from_row(obj, row)
            else:
                obj = HsRate(
                    hs_code=pref,
                    hs_prefix=pref,
                    duty_rate=duty_stored if duty_stored is not None else "0",
                    source_revision=rev[:128],
                    source_url="bulk-normative-ai",
                )
                _apply_hs_rate_fields_from_row(obj, row)
                db.add(obj)
            applied += 1

        elif cat == "excise" and prefix_union:
            pct = _parse_percent(rate_txt)
            fixed = _parse_fixed_rub(rate_txt)
            pref = prefix_union[:10]
            obj = db.query(HsRate).filter(HsRate.hs_prefix == pref).first()
            if fixed is not None and (pct is None or ("руб" in rate_txt.lower() or "фикс" in rate_txt.lower())):
                ex_type = "fixed"
                ex_val = float(fixed)
                basis = rate_txt[:500] or act
            elif pct is not None:
                ex_type = "percent"
                ex_val = float(pct)
                basis = rate_txt[:500] or act
            else:
                continue
            if obj:
                obj.excise_type = ex_type
                obj.excise_value = ex_val
                obj.excise_basis = basis[:4000]
                obj.source_revision = rev[:128]
                _apply_hs_rate_fields_from_row(obj, row)
            else:
                obj = HsRate(
                    hs_code=pref,
                    hs_prefix=pref,
                    duty_rate="0",
                    excise_type=ex_type,
                    excise_value=ex_val,
                    excise_basis=basis[:4000],
                    source_revision=rev[:128],
                    source_url="bulk-normative-ai",
                )
                _apply_hs_rate_fields_from_row(obj, row)
                db.add(obj)
            applied += 1

    return applied


async def call_gemini_with_throttle(
    document_text: str,
    *,
    min_interval_sec: float = 4.0,
    max_retries: int = 8,
) -> str:
    """Вызов Gemini с паузой между запросами и экспоненциальным backoff при 429."""
    global _last_llm_monotonic

    from pathlib import Path

    from dotenv import load_dotenv
    from google.api_core.client_options import ClientOptions

    _backend_root = Path(__file__).resolve().parents[2]
    load_dotenv(_backend_root / ".env")
    load_dotenv()

    key = (os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY") or "").strip()
    if not key:
        raise RuntimeError("Не задан GEMINI_API_KEY или GOOGLE_API_KEY")

    try:
        import google.generativeai as genai
    except ModuleNotFoundError as e:
        raise RuntimeError("Не установлен пакет google-generativeai") from e

    model_name = resolved_gemini_model_name()
    base_url = normalize_gemini_api_endpoint_for_sdk(os.getenv("GEMINI_BASE_URL") or "")
    if base_url:
        genai.configure(
            api_key=key,
            transport="rest",
            client_options=ClientOptions(api_endpoint=base_url),
        )
    else:
        genai.configure(api_key=key)
    model = genai.GenerativeModel(model_name, system_instruction=BULK_SYSTEM_PROMPT)

    body = (document_text or "").strip()
    if len(body) > _MAX_TEXT_CHARS:
        body = body[:_MAX_TEXT_CHARS] + "\n\n[…текст обрезан для лимита модели…]"

    attempt = 0
    backoff_base = 60.0

    while True:
        now = time.monotonic()
        wait = min_interval_sec - (now - _last_llm_monotonic)
        if wait > 0:
            await asyncio.sleep(wait)

        def _call() -> str:
            resp = model.generate_content(
                "Проанализируй следующий фрагмент нормативного текста и верни только JSON-массив по инструкции.\n\n"
                + body,
                generation_config={"temperature": 0.15, "max_output_tokens": 8192},
            )
            return (getattr(resp, "text", "") or "").strip()

        try:
            raw = await asyncio.to_thread(_call)
            _last_llm_monotonic = time.monotonic()
            _save_llm_raw_response(raw)
            return raw
        except Exception as e:
            _last_llm_monotonic = time.monotonic()
            if _is_rate_limit_error(e) and attempt < max_retries:
                pause = min(backoff_base * (2**attempt), 900.0)
                logger.warning(f"Gemini 429/rate-limit, пауза {pause:.0f} с (попытка {attempt + 1}/{max_retries}): {e}")
                await asyncio.sleep(pause)
                attempt += 1
                continue
            raise


def parse_llm_json_array(raw: str) -> list[dict[str, Any]]:
    payload = _extract_json_array(raw)
    data = json.loads(payload)
    if not isinstance(data, list):
        return []
    out: list[dict[str, Any]] = []
    for el in data:
        if isinstance(el, dict):
            out.append(el)
    return out


def _checkpoint_ok(db: Session, digest: str) -> bool:
    row = db.query(BulkImportFileCheckpoint).filter(BulkImportFileCheckpoint.file_sha256 == digest).first()
    return row is not None and (row.status or "") == "ok"


def _save_checkpoint(
    db: Session,
    *,
    digest: str,
    rel: str,
    status: str,
    measures: int,
    err: str,
    job_id: int | None,
) -> None:
    row = db.query(BulkImportFileCheckpoint).filter(BulkImportFileCheckpoint.file_sha256 == digest).first()
    now = utc_now_naive()
    if row is None:
        db.add(
            BulkImportFileCheckpoint(
                file_sha256=digest,
                relative_path=rel[:512],
                status=status,
                measures_applied=measures,
                error_note=err[:4000],
                job_id=job_id,
                processed_at=now,
            )
        )
    else:
        row.relative_path = rel[:512]
        row.status = status
        row.measures_applied = measures
        row.error_note = err[:4000]
        row.job_id = job_id
        row.processed_at = now


def _update_job(
    db: Session,
    job_id: int,
    *,
    status: str | None = None,
    processed: int | None = None,
    total: int | None = None,
    measures: int | None = None,
    current: str | None = None,
    error: str | None = None,
) -> None:
    job = db.query(BulkImportJob).filter(BulkImportJob.id == job_id).first()
    if not job:
        return
    if status is not None:
        job.status = status
    if processed is not None:
        job.processed_files = processed
    if total is not None:
        job.total_files = total
    if measures is not None:
        job.measures_applied = measures
    if current is not None:
        job.current_file = current[:512]
    if error is not None:
        job.error_message = error[:4000]
    job.updated_at = utc_now_naive()


async def run_bulk_import(
    job_id: int,
    *,
    delay_sec: float = 4.0,
    skip_checkpoint: bool = False,
    progress_cb: Callable[[dict[str, Any]], None] | None = None,
) -> None:
    """
    Обрабатывает все файлы в data/raw_normative. Идемпотентность по SHA-256 в bulk_import_file_checkpoints.
    """
    global _job_running
    async with _import_lock:
        if _job_running:
            raise RuntimeError("Уже выполняется другая задача импорта")
        _job_running = True
        try:
            root = raw_normative_dir()
            files = list_input_files(root)
            with SessionLocal() as db:
                _update_job(db, job_id, total=len(files), status="running", processed=0, current="", error="")
                db.commit()

            total_measures = 0
            done = 0
            for path in files:
                rel = str(path.relative_to(root))
                digest = _sha256_file(path)
                with SessionLocal() as db:
                    if not skip_checkpoint and _checkpoint_ok(db, digest):
                        done += 1
                        _update_job(db, job_id, processed=done, current=rel)
                        db.commit()
                        if progress_cb:
                            progress_cb(
                                {
                                    "processed_files": done,
                                    "total_files": len(files),
                                    "measures_applied": total_measures,
                                    "skipped": True,
                                    "file": rel,
                                }
                            )
                        continue

                try:
                    text = extract_text_from_file(path)
                except Exception as e:
                    logger.warning(f"bulk import: не удалось прочитать {rel}: {e}")
                    with SessionLocal() as db:
                        _save_checkpoint(
                            db, digest=digest, rel=rel, status="error", measures=0, err=str(e), job_id=job_id
                        )
                        done += 1
                        _update_job(db, job_id, processed=done, current=rel, error=str(e))
                        db.commit()
                    if progress_cb:
                        progress_cb(
                            {
                                "processed_files": done,
                                "total_files": len(files),
                                "measures_applied": total_measures,
                                "error": str(e),
                                "file": rel,
                            }
                        )
                    await asyncio.sleep(delay_sec)
                    continue

                if not text.strip():
                    with SessionLocal() as db:
                        _save_checkpoint(
                            db,
                            digest=digest,
                            rel=rel,
                            status="error",
                            measures=0,
                            err="Пустой текст после извлечения",
                            job_id=job_id,
                        )
                        done += 1
                        _update_job(db, job_id, processed=done, current=rel)
                        db.commit()
                    await asyncio.sleep(delay_sec)
                    continue

                try:
                    raw = await call_gemini_with_throttle(text, min_interval_sec=delay_sec)
                    rows = parse_llm_json_array(raw)
                except Exception as e:
                    logger.exception(f"bulk import LLM {rel}: {e}")
                    with SessionLocal() as db:
                        _save_checkpoint(
                            db, digest=digest, rel=rel, status="error", measures=0, err=str(e), job_id=job_id
                        )
                        done += 1
                        _update_job(db, job_id, processed=done, current=rel, error=str(e))
                        db.commit()
                    if progress_cb:
                        progress_cb(
                            {
                                "processed_files": done,
                                "total_files": len(files),
                                "measures_applied": total_measures,
                                "error": str(e),
                                "file": rel,
                            }
                        )
                    await asyncio.sleep(delay_sec)
                    continue

                with SessionLocal() as db:
                    try:
                        n = apply_structured_rows(db, rows, source_tag=rel)
                        db.commit()
                    except Exception as e:
                        db.rollback()
                        logger.exception(f"bulk import DB {rel}: {e}")
                        _save_checkpoint(
                            db, digest=digest, rel=rel, status="error", measures=0, err=str(e), job_id=job_id
                        )
                        done += 1
                        _update_job(db, job_id, processed=done, current=rel, error=str(e))
                        db.commit()
                        if progress_cb:
                            progress_cb(
                                {
                                    "processed_files": done,
                                    "total_files": len(files),
                                    "measures_applied": total_measures,
                                    "error": str(e),
                                    "file": rel,
                                }
                            )
                        await asyncio.sleep(delay_sec)
                        continue

                    _save_checkpoint(db, digest=digest, rel=rel, status="ok", measures=n, err="", job_id=job_id)
                    total_measures += n
                    done += 1
                    _update_job(db, job_id, processed=done, measures=total_measures, current=rel, error="")
                    db.commit()

                if progress_cb:
                    progress_cb(
                        {
                            "processed_files": done,
                            "total_files": len(files),
                            "measures_applied": total_measures,
                            "file": rel,
                            "llm_rows": len(rows),
                        }
                    )

                await asyncio.sleep(delay_sec)

            with SessionLocal() as db:
                _update_job(
                    db, job_id, status="completed", processed=done, measures=total_measures, current="", error=""
                )
                db.commit()
            try:
                from .preview_cache_revision import bump_preview_cache_revision

                if total_measures > 0:
                    bump_preview_cache_revision("bulk_normative_ai")
            except Exception as e:
                logger.warning(f"bump_preview_cache_revision: {e}")
        except Exception as e:
            logger.exception(f"bulk import job {job_id}: {e}")
            with SessionLocal() as db:
                _update_job(db, job_id, status="error", error=str(e))
                db.commit()
            raise
        finally:
            _job_running = False


def get_job_status(job_id: int | None = None) -> dict[str, Any]:
    with SessionLocal() as db:
        if job_id is not None:
            job = db.query(BulkImportJob).filter(BulkImportJob.id == job_id).first()
        else:
            job = db.query(BulkImportJob).order_by(BulkImportJob.id.desc()).first()
        if not job:
            return {"job": None}
        return {
            "job": {
                "id": job.id,
                "status": job.status,
                "total_files": job.total_files,
                "processed_files": job.processed_files,
                "measures_applied": job.measures_applied,
                "current_file": job.current_file,
                "error_message": job.error_message,
                "created_at": job.created_at.isoformat() if job.created_at else None,
                "updated_at": job.updated_at.isoformat() if job.updated_at else None,
            },
            "worker_busy": is_import_running(),
        }


def get_latest_job_status() -> dict[str, Any]:
    return get_job_status(None)


def create_import_job() -> int:
    with SessionLocal() as db:
        j = BulkImportJob(
            status="queued",
            total_files=0,
            processed_files=0,
            measures_applied=0,
            current_file="",
            error_message="",
        )
        db.add(j)
        db.commit()
        db.refresh(j)
        return int(j.id)


def is_import_running() -> bool:
    return _job_running
